import requests
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE

# Create an instance of a word document
doc = docx.Document()

style = doc.styles
new_heading_style = style.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
new_heading_style.base_style = style['Title']
font = new_heading_style.font
font.name = 'Comic Sans MS'
font.size = Pt(20)
font.bold = True

# For h1
nhs1 = style.add_style('nhs1', WD_STYLE_TYPE.PARAGRAPH)
nhs1.base_style = style['Heading 1']
font = nhs1.font
font.name = 'Comic Sans MS'
font.size = Pt(14)

# For h2
nhs2 = style.add_style('nhs2', WD_STYLE_TYPE.PARAGRAPH)
nhs2.base_style = style['Heading 2']
font = nhs2.font
font.name = 'Comic Sans MS'
font.size = Pt(13)

u = input("Enter key word: ")
url = f'https://en.wikipedia.org/w/index.php?search={u}&title=Special:Search&profile=advanced&fulltext=1&ns0=1'

r = requests.get(url)
htmlContent = r.content

soup = BeautifulSoup(htmlContent, 'html.parser')

link_div = soup.find_all('div',attrs={'class':'mw-search-result-heading'})[0]
link = link_div.find_all('a')[0]
link = "https://en.wikipedia.org" + link.get('href')

r = requests.get(link)
htmlContent = r.content

soup = BeautifulSoup(htmlContent, 'html.parser')

fileName = list(link.split("/"))[-1]

s = soup.find_all(['sup', 'img', 'math'])
for s1 in s:
    s1.clear()

edits = soup.find_all('span',attrs={'class':'mw-editsection'})
for edit in edits:
    edit.clear()

elems = soup.find_all(['p', 'h1', 'h2', 'h3'])
for elem in elems:
    if elem.name == 'p':
        P1 = doc.add_paragraph()
        P1.paragraph_format.first_line_indent = Inches(0.5)
        P = P1.add_run(elem.get_text())
        P.font.size = Pt(12)
        P.font.name = 'Roboto'
    elif elem.name == 'h1':
        H = doc.add_paragraph(elem.get_text().upper(), style='New Heading')
    elif elem.name == 'h2':
        doc.add_paragraph(elem.get_text().upper(), style='nhs1')
    else:
        doc.add_paragraph(elem.get_text().upper(), style='nhs2')

doc.save(f"{fileName}.docx")

